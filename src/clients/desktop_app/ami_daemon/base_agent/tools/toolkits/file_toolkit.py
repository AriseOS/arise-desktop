"""
FileToolkit - File writing and management for document operations.

Based on Eigent's FileToolkit implementation which wraps CAMEL's FileToolkit.
Provides file writing capabilities with event emission support.

References:
- Eigent: third-party/eigent/backend/app/utils/toolkit/file_write_toolkit.py
- CAMEL: camel.toolkits.FileToolkit
"""

import logging
from pathlib import Path
from typing import List, Optional, Union

from .base_toolkit import BaseToolkit, FunctionTool
from ...events import listen_toolkit
from ...workspace import get_working_directory

logger = logging.getLogger(__name__)


class FileToolkit(BaseToolkit):
    """A toolkit for file writing and management operations.

    Provides file creation and writing capabilities:
    - Write content to various file formats (txt, md, html, json, csv, etc.)
    - UTF-8 encoding by default
    - Automatic backup functionality
    - Support for LaTeX rendering in PDF

    Based on Eigent's implementation which wraps CAMEL's FileToolkit.
    """

    agent_name: str = "document_agent"

    def __init__(
        self,
        working_directory: Optional[str] = None,
        timeout: Optional[float] = 60.0,
        default_encoding: str = "utf-8",
        backup_enabled: bool = True,
    ) -> None:
        """Initialize the FileToolkit.

        Args:
            working_directory: Directory for file operations.
                If not provided, uses task workspace from WorkingDirectoryManager.
            timeout: Operation timeout in seconds.
            default_encoding: Default file encoding (default: utf-8).
            backup_enabled: Whether to create backups when modifying files.
        """
        super().__init__(timeout=timeout)

        # Determine working directory - fail if not provided and no workspace manager
        if working_directory:
            self._working_directory = Path(working_directory)
        else:
            self._working_directory = Path(get_working_directory())

        # Ensure directory exists
        self._working_directory.mkdir(parents=True, exist_ok=True)

        self._default_encoding = default_encoding
        self._backup_enabled = backup_enabled

        logger.info(
            f"FileToolkit initialized in {self._working_directory} "
            f"(encoding={default_encoding}, backup={backup_enabled})"
        )

    def _resolve_filepath(self, filename: str) -> Path:
        """Resolve a filename to an absolute path.

        If the filename is absolute, returns it directly.
        Otherwise, joins it with the working directory.

        Args:
            filename: The filename or path.

        Returns:
            Absolute Path object.
        """
        path = Path(filename)
        if path.is_absolute():
            return path
        return self._working_directory / filename

    def _create_backup(self, filepath: Path) -> Optional[Path]:
        """Create a backup of a file if it exists.

        Args:
            filepath: Path to the file.

        Returns:
            Path to backup file if created, None otherwise.
        """
        if not self._backup_enabled or not filepath.exists():
            return None

        import shutil
        import datetime

        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        backup_path = filepath.with_suffix(f".{timestamp}.bak{filepath.suffix}")

        try:
            shutil.copy2(filepath, backup_path)
            logger.info(f"Created backup: {backup_path}")
            return backup_path
        except Exception as e:
            logger.warning(f"Failed to create backup: {e}")
            return None

    def _markdown_to_html(self, content: str) -> str:
        """Convert Markdown content to HTML.

        Shared by DOCX and PDF paths.

        Args:
            content: Markdown-formatted string.

        Returns:
            HTML string.
        """
        import markdown

        return markdown.markdown(
            content,
            extensions=["tables", "fenced_code", "sane_lists"],
        )

    def _register_cjk_font(self) -> str:
        """Register a CJK-capable font for PDF generation.

        Searches for platform-specific font files and registers the first
        available one with reportlab. Falls back to Helvetica.

        Returns:
            Registered font name for use in reportlab styles.
        """
        import os
        import platform as plat

        from reportlab.lib.fonts import addMapping
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont

        font_paths = []
        system = plat.system()

        if system == "Darwin":
            font_paths = [
                "/System/Library/Fonts/STHeiti Light.ttc",
                "/System/Library/Fonts/STHeiti Medium.ttc",
                "/System/Library/Fonts/Supplemental/Songti.ttc",
                "/System/Library/Fonts/PingFang.ttc",
                "/Library/Fonts/Arial Unicode.ttf",
            ]
        elif system == "Windows":
            # Prefer .ttf over .ttc (fewer postscript outline issues)
            font_paths = [
                r"C:\Windows\Fonts\msyh.ttf",
                r"C:\Windows\Fonts\msyh.ttc",
                r"C:\Windows\Fonts\simsun.ttf",
                r"C:\Windows\Fonts\simsun.ttc",
                r"C:\Windows\Fonts\simhei.ttf",
                r"C:\Windows\Fonts\malgun.ttf",
            ]
        elif system == "Linux":
            font_paths = [
                "/usr/share/fonts/truetype/wqy/wqy-microhei.ttc",
                "/usr/share/fonts/truetype/droid/DroidSansFallbackFull.ttf",
                "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
                "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
            ]

        for font_path in font_paths:
            if os.path.exists(font_path):
                try:
                    font_name = "CJKFont"
                    if font_name not in pdfmetrics.getRegisteredFontNames():
                        pdfmetrics.registerFont(TTFont(font_name, font_path))
                        addMapping(font_name, 0, 0, font_name)
                        addMapping(font_name, 0, 1, font_name)
                        addMapping(font_name, 1, 0, font_name)
                        addMapping(font_name, 1, 1, font_name)
                    return font_name
                except Exception:
                    continue

        logger.warning("No CJK font found, falling back to Helvetica")
        return "Helvetica"

    def _html_to_pdf_flowables(self, html: str, font_name: str) -> list:
        """Convert HTML string to a list of reportlab flowables.

        Parses HTML block elements and maps them to reportlab equivalents.

        Args:
            html: HTML string (from _markdown_to_html).
            font_name: Font name to use in paragraph styles.

        Returns:
            List of reportlab flowable objects.
        """
        from bs4 import BeautifulSoup
        from reportlab.lib import colors
        from reportlab.lib.styles import ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import (
            HRFlowable,
            ListFlowable,
            ListItem,
            Paragraph,
            Preformatted,
            Spacer,
            Table,
            TableStyle,
        )

        # Define styles
        styles = {
            "body": ParagraphStyle(
                "body", fontName=font_name, fontSize=11, leading=15,
                spaceAfter=8,
            ),
            "h1": ParagraphStyle(
                "h1", fontName=font_name, fontSize=22, leading=28,
                spaceAfter=12, spaceBefore=16, textColor=colors.HexColor("#1a1a1a"),
            ),
            "h2": ParagraphStyle(
                "h2", fontName=font_name, fontSize=18, leading=24,
                spaceAfter=10, spaceBefore=14, textColor=colors.HexColor("#1a1a1a"),
            ),
            "h3": ParagraphStyle(
                "h3", fontName=font_name, fontSize=15, leading=20,
                spaceAfter=8, spaceBefore=12, textColor=colors.HexColor("#333333"),
            ),
            "h4": ParagraphStyle(
                "h4", fontName=font_name, fontSize=13, leading=18,
                spaceAfter=6, spaceBefore=10, textColor=colors.HexColor("#333333"),
            ),
            "code": ParagraphStyle(
                "code", fontName="Courier", fontSize=9, leading=12,
                spaceAfter=8, leftIndent=12, rightIndent=12,
                backColor=colors.HexColor("#f5f5f5"),
            ),
            "blockquote": ParagraphStyle(
                "blockquote", fontName=font_name, fontSize=11, leading=15,
                leftIndent=24, spaceAfter=8, textColor=colors.HexColor("#555555"),
            ),
        }

        soup = BeautifulSoup(html, "html.parser")
        flowables = []

        for element in soup.children:
            if element.name is None:
                # NavigableString (whitespace)
                text = str(element).strip()
                if text:
                    flowables.append(Paragraph(text, styles["body"]))
                continue

            tag = element.name

            if tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
                level = tag  # e.g. "h1"
                style = styles.get(level, styles["h4"])
                flowables.append(Paragraph(str(element.decode_contents()), style))

            elif tag == "p":
                flowables.append(
                    Paragraph(str(element.decode_contents()), styles["body"])
                )

            elif tag in ("ul", "ol"):
                items = []
                for li in element.find_all("li", recursive=False):
                    items.append(
                        ListItem(
                            Paragraph(str(li.decode_contents()), styles["body"])
                        )
                    )
                if items:
                    if tag == "ul":
                        flowables.append(
                            ListFlowable(items, bulletType="bullet")
                        )
                    else:
                        flowables.append(
                            ListFlowable(
                                items,
                                bulletType="1",
                                bulletFontName="Helvetica",
                                start=1,
                            )
                        )

            elif tag == "table":
                rows_data = []
                for tr in element.find_all("tr"):
                    cells = tr.find_all(["th", "td"])
                    rows_data.append(
                        [Paragraph(str(c.decode_contents()), styles["body"]) for c in cells]
                    )
                if rows_data:
                    table_style_cmds = [
                        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                        ("TOPPADDING", (0, 0), (-1, -1), 4),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                        ("LEFTPADDING", (0, 0), (-1, -1), 6),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                    ]
                    # Grey background for header row
                    if element.find("th"):
                        table_style_cmds.append(
                            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#e0e0e0"))
                        )
                    t = Table(rows_data)
                    t.setStyle(TableStyle(table_style_cmds))
                    flowables.append(t)
                    flowables.append(Spacer(1, 8))

            elif tag == "pre":
                code_tag = element.find("code")
                code_text = code_tag.get_text() if code_tag else element.get_text()
                flowables.append(Preformatted(code_text, styles["code"]))

            elif tag == "blockquote":
                inner = element.get_text(separator="\n").strip()
                flowables.append(Paragraph(inner, styles["blockquote"]))

            elif tag == "hr":
                flowables.append(HRFlowable(width="100%", color=colors.grey))

            else:
                # Fallback: render as paragraph
                text = str(element.decode_contents()).strip()
                if text:
                    flowables.append(Paragraph(text, styles["body"]))

        return flowables

    @listen_toolkit(
        inputs=lambda self, title, content, filename, **kw: f"Writing to {filename}",
        return_msg=lambda r: r[:200] if isinstance(r, str) else str(r)
    )
    def write_to_file(
        self,
        title: str,
        content: Union[str, List[List[str]]],
        filename: str,
        encoding: Optional[str] = None,
        use_latex: bool = False,
    ) -> str:
        """Write content to a file.

        Supports various file formats including:
        - Text files (.txt, .md, .html, .json, .yaml, .xml)
        - CSV files (.csv) - content can be a list of lists
        - Word documents (.docx) - Markdown auto-converted to formatted Word
        - PDF files (.pdf) - Markdown auto-converted to formatted PDF with CJK support

        For .docx and .pdf, write content in Markdown. Supported elements:
        headings (#-####), bold, italic, inline code, bullet/numbered lists,
        tables, fenced code blocks, blockquotes, and horizontal rules.

        Args:
            title: Title or heading for the content.
            content: The content to write. For CSV, can be a list of lists.
                For .docx and .pdf, use Markdown formatting.
            filename: The filename to write to.
            encoding: File encoding (default: utf-8).
            use_latex: Whether to render LaTeX in PDF files.

        Returns:
            Success message with file path, or error message.
        """
        effective_encoding = encoding or self._default_encoding
        filepath = self._resolve_filepath(filename)

        # Create parent directories if needed
        filepath.parent.mkdir(parents=True, exist_ok=True)

        # Create backup if file exists
        self._create_backup(filepath)

        logger.info(f"Writing to file: {filepath}")

        try:
            suffix = filepath.suffix.lower()

            if suffix == ".csv":
                # Handle CSV files
                import csv
                with open(filepath, "w", newline="", encoding=effective_encoding) as f:
                    writer = csv.writer(f)
                    if isinstance(content, list):
                        for row in content:
                            writer.writerow(row)
                    else:
                        # Write content as single column
                        for line in content.split("\n"):
                            writer.writerow([line])

            elif suffix == ".json":
                # Handle JSON files
                import json
                with open(filepath, "w", encoding=effective_encoding) as f:
                    if isinstance(content, str):
                        # Try to parse and pretty-print
                        try:
                            data = json.loads(content)
                            json.dump(data, f, indent=2, ensure_ascii=False)
                        except json.JSONDecodeError:
                            f.write(content)
                    else:
                        json.dump(content, f, indent=2, ensure_ascii=False)

            elif suffix == ".docx":
                # Handle Word documents (Markdown → HTML → DOCX)
                try:
                    from docx import Document
                    from htmldocx import HtmlToDocx

                    doc = Document()

                    # Add title
                    if title:
                        doc.add_heading(title, level=1)

                    # Convert content
                    if isinstance(content, str):
                        html = self._markdown_to_html(content)
                        parser = HtmlToDocx()
                        parser.add_html_to_document(html, doc)
                    elif isinstance(content, list):
                        for row in content:
                            doc.add_paragraph(", ".join(str(cell) for cell in row))

                    doc.save(str(filepath))
                except ImportError:
                    return "Error: python-docx and htmldocx packages required for .docx files. Install with: pip install python-docx htmldocx"

            elif suffix == ".pdf":
                # Handle PDF files (Markdown → HTML → reportlab flowables)
                try:
                    from reportlab.lib.pagesizes import letter
                    from reportlab.lib.styles import ParagraphStyle
                    from reportlab.lib.units import inch
                    from reportlab.platypus import (
                        Paragraph,
                        SimpleDocTemplate,
                        Spacer,
                    )

                    font_name = self._register_cjk_font()

                    doc = SimpleDocTemplate(
                        str(filepath),
                        pagesize=letter,
                        leftMargin=1 * inch,
                        rightMargin=1 * inch,
                        topMargin=1 * inch,
                        bottomMargin=1 * inch,
                    )

                    flowables = []

                    # Add title
                    if title:
                        title_style = ParagraphStyle(
                            "title",
                            fontName=font_name,
                            fontSize=24,
                            leading=30,
                            spaceAfter=16,
                        )
                        flowables.append(Paragraph(title, title_style))
                        flowables.append(Spacer(1, 8))

                    # Convert content
                    if isinstance(content, str):
                        html = self._markdown_to_html(content)
                        flowables.extend(self._html_to_pdf_flowables(html, font_name))
                    elif isinstance(content, list):
                        body_style = ParagraphStyle(
                            "body_list", fontName=font_name, fontSize=11, leading=15,
                        )
                        for row in content:
                            flowables.append(
                                Paragraph(", ".join(str(cell) for cell in row), body_style)
                            )

                    doc.build(flowables)
                except ImportError:
                    return "Error: reportlab package required for .pdf files. Install with: pip install reportlab"

            else:
                # Default: write as text file
                full_content = ""
                if title:
                    # Add title based on file type
                    if suffix == ".md":
                        full_content = f"# {title}\n\n"
                    elif suffix == ".html":
                        full_content = f"<h1>{title}</h1>\n\n"
                    else:
                        full_content = f"{title}\n{'=' * len(title)}\n\n"

                if isinstance(content, str):
                    full_content += content
                elif isinstance(content, list):
                    # Convert list of lists to text
                    for row in content:
                        full_content += "\t".join(str(cell) for cell in row) + "\n"

                with open(filepath, "w", encoding=effective_encoding) as f:
                    f.write(full_content)

            logger.info(f"Successfully wrote to file: {filepath}")
            return f"Content successfully written to file: {filepath}"

        except Exception as e:
            error_msg = f"Error writing to file {filepath}: {str(e)}"
            logger.error(error_msg)
            return error_msg

    @listen_toolkit(
        inputs=lambda self, content, filename, **kw: f"Appending to {filename}",
        return_msg=lambda r: r[:200] if isinstance(r, str) else str(r)
    )
    def append_to_file(
        self,
        content: str,
        filename: str,
        encoding: Optional[str] = None,
    ) -> str:
        """Append content to an existing file. Creates the file if it does not exist.

        Use this tool instead of write_to_file when you need to ADD data to a file
        without losing existing content. Ideal for:
        - .jsonl files: each call appends one JSON object per line
        - .txt/.md files: each call appends text to the end
        - .csv files: each call appends rows to the end
        - Any text-based file that accumulates data across multiple steps

        Args:
            content: The content to append. For .jsonl files, should be a valid JSON string.
            filename: The filename to append to.
            encoding: File encoding (default: utf-8).

        Returns:
            Success message with file path, or error message.
        """
        effective_encoding = encoding or self._default_encoding
        filepath = self._resolve_filepath(filename)

        filepath.parent.mkdir(parents=True, exist_ok=True)

        logger.info(f"Appending to file: {filepath}")

        try:
            suffix = filepath.suffix.lower()

            if suffix == ".jsonl":
                import json
                # Validate and normalize JSON, write as single line
                try:
                    data = json.loads(content)
                    line = json.dumps(data, ensure_ascii=False) + "\n"
                except json.JSONDecodeError:
                    line = content.rstrip("\n") + "\n"

                with open(filepath, "a", encoding=effective_encoding) as f:
                    f.write(line)
            else:
                with open(filepath, "a", encoding=effective_encoding) as f:
                    f.write(content)

            logger.info(f"Successfully appended to file: {filepath}")
            return f"Content successfully appended to file: {filepath}"

        except Exception as e:
            error_msg = f"Error appending to file {filepath}: {str(e)}"
            logger.error(error_msg)
            return error_msg

    @listen_toolkit(
        inputs=lambda self, filepath: f"Reading {filepath}",
        return_msg=lambda r: f"Read {len(r)} chars" if isinstance(r, str) else str(r)
    )
    def read_file(self, filepath: str) -> str:
        """Read content from a file.

        Args:
            filepath: Path to the file to read.

        Returns:
            File content as string, or error message.
        """
        path = self._resolve_filepath(filepath)

        if not path.exists():
            return f"Error: File not found: {path}"

        try:
            with open(path, "r", encoding=self._default_encoding) as f:
                content = f.read()
            logger.info(f"Read {len(content)} chars from {path}")
            return content
        except Exception as e:
            error_msg = f"Error reading file {path}: {str(e)}"
            logger.error(error_msg)
            return error_msg

    @listen_toolkit(
        inputs=lambda self, filepath: f"Checking {filepath}",
        return_msg=lambda r: str(r)
    )
    def file_exists(self, filepath: str) -> bool:
        """Check if a file exists.

        Args:
            filepath: Path to check.

        Returns:
            True if file exists, False otherwise.
        """
        path = self._resolve_filepath(filepath)
        return path.exists()

    @listen_toolkit(
        inputs=lambda self, directory=None: f"Listing files in {directory or 'working directory'}",
        return_msg=lambda r: f"Found {len(r)} files" if isinstance(r, list) else str(r)
    )
    def list_files(self, directory: Optional[str] = None) -> List[str]:
        """List files in a directory.

        Args:
            directory: Directory to list. Defaults to working directory.

        Returns:
            List of filenames.
        """
        dir_path = self._resolve_filepath(directory) if directory else self._working_directory

        if not dir_path.exists():
            logger.warning(f"Directory not found: {dir_path}")
            return []

        try:
            files = [f.name for f in dir_path.iterdir() if f.is_file()]
            return sorted(files)
        except Exception as e:
            logger.error(f"Error listing files: {e}")
            return []

    @property
    def working_directory(self) -> Path:
        """Get the current working directory."""
        return self._working_directory

    def get_tools(self) -> List[FunctionTool]:
        """Return a list of FunctionTool objects for this toolkit.

        Returns:
            List of FunctionTool objects.
        """
        return [
            FunctionTool(self.write_to_file),
            FunctionTool(self.append_to_file),
            FunctionTool(self.read_file),
            FunctionTool(self.file_exists),
            FunctionTool(self.list_files),
        ]

    @classmethod
    def toolkit_name(cls) -> str:
        """Return the name of this toolkit."""
        return "File"
